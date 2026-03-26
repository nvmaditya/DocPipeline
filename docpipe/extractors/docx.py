from __future__ import annotations

from typing import Any, Dict, List, Optional

from docx import Document

from .base import BaseExtractor


class DocxExtractor(BaseExtractor):
    def extract(self, file_path: str) -> List[Dict[str, Any]]:
        doc = Document(file_path)
        records: List[Dict[str, Any]] = []
        current_heading: Optional[str] = None

        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.lower().startswith("heading"):
                current_heading = text
            records.append(
                {
                    "text": text,
                    "page_number": None,
                    "heading_context": current_heading,
                }
            )

        for table in doc.tables:
            rows: List[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(r for r in rows if r.strip())
            if table_text:
                records.append(
                    {
                        "text": table_text,
                        "page_number": None,
                        "heading_context": current_heading,
                    }
                )

        return records
